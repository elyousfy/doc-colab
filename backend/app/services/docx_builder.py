"""Build a .docx file from Tiptap-compatible JSON content."""

from __future__ import annotations

import io
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def build_docx(content: dict, images: dict[str, bytes]) -> bytes:
    """Convert Tiptap JSON block schema to DOCX bytes.

    images maps src_url → image_bytes.
    """
    doc = Document()
    nodes = content.get("content", [])

    header_nodes = []
    footer_nodes = []
    body_nodes = []

    for node in nodes:
        if node.get("type") == "docSection":
            section_type = (node.get("attrs") or {}).get("sectionType", "")
            if section_type == "header":
                header_nodes = node.get("content", [])
                continue
            elif section_type == "footer":
                footer_nodes = node.get("content", [])
                continue
        body_nodes.append(node)

    if header_nodes:
        section = doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        # Clear default empty paragraph
        for p in header.paragraphs:
            p.clear()
        _render_nodes_to_container(header, header_nodes, images, is_first=True)

    for node in body_nodes:
        _render_node(doc, node, images)

    if footer_nodes:
        section = doc.sections[0]
        footer = section.footer
        footer.is_linked_to_previous = False
        for p in footer.paragraphs:
            p.clear()
        _render_nodes_to_container(footer, footer_nodes, images, is_first=True)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _render_nodes_to_container(
    container, nodes: list[dict], images: dict[str, bytes], is_first: bool = False
) -> None:
    """Render nodes into a header/footer container."""
    for i, node in enumerate(nodes):
        node_type = node.get("type", "")
        attrs = node.get("attrs") or {}
        content = node.get("content", [])

        if i == 0 and is_first and container.paragraphs:
            para = container.paragraphs[0]
        else:
            para = container.add_paragraph()

        if node_type == "image":
            src = attrs.get("src", "")
            img_data = images.get(src)
            if img_data:
                width = attrs.get("width")
                run = para.add_run()
                run.add_picture(io.BytesIO(img_data), width=Pt(width) if width else None)
        elif node_type in ("paragraph", "heading"):
            _apply_alignment(para, attrs)
            _apply_inline_content(para, content, images)
        elif node_type == "docSection":
            for child in content:
                child_type = child.get("type", "")
                child_attrs = child.get("attrs") or {}
                child_content = child.get("content", [])
                p = container.add_paragraph()
                _apply_alignment(p, child_attrs)
                _apply_inline_content(p, child_content, images)


# ---------------------------------------------------------------------------
# Node dispatcher
# ---------------------------------------------------------------------------

def _render_node(doc: Document, node: dict, images: dict[str, bytes]) -> None:
    node_type = node.get("type", "")
    attrs = node.get("attrs", {})

    if node_type == "heading":
        _render_heading(doc, node, attrs, images)
    elif node_type == "paragraph":
        _render_paragraph(doc, node, attrs, images)
    elif node_type == "image":
        _render_image(doc, node, attrs, images)
    elif node_type == "table":
        _render_table(doc, node, images)
    elif node_type == "bulletList":
        _render_list(doc, node, images, style="List Bullet")
    elif node_type == "orderedList":
        _render_list(doc, node, images, style="List Number")
    elif node_type == "blockquote":
        _render_blockquote(doc, node, images)
    elif node_type == "hardBreak":
        if attrs.get("pageBreak"):
            _render_page_break(doc)
    elif node_type == "horizontalRule":
        _render_horizontal_rule(doc)
    elif node_type == "docSection":
        _render_doc_section(doc, node, images)


# ---------------------------------------------------------------------------
# DocSection (TOC in body — header/footer handled above)
# ---------------------------------------------------------------------------

def _render_doc_section(
    doc: Document, node: dict, images: dict[str, bytes]
) -> None:
    """Render a docSection (e.g. TOC) as body content."""
    for child in node.get("content", []):
        _render_node(doc, child, images)


# ---------------------------------------------------------------------------
# Heading
# ---------------------------------------------------------------------------

def _render_heading(
    doc: Document, node: dict, attrs: dict, images: dict[str, bytes]
) -> None:
    level = attrs.get("level", 1)
    level = max(1, min(level, 9))
    heading = doc.add_heading("", level=level)
    _apply_alignment(heading, attrs)
    _apply_inline_content(heading, node.get("content", []), images)


# ---------------------------------------------------------------------------
# Paragraph
# ---------------------------------------------------------------------------

def _render_paragraph(
    doc: Document, node: dict, attrs: dict, images: dict[str, bytes]
) -> None:
    para = doc.add_paragraph()
    _apply_alignment(para, attrs)
    _apply_inline_content(para, node.get("content", []), images)


# ---------------------------------------------------------------------------
# Image
# ---------------------------------------------------------------------------

def _render_image(
    doc: Document, node: dict, attrs: dict, images: dict[str, bytes]
) -> None:
    src = attrs.get("src", "")
    img_data = images.get(src)
    if not img_data:
        return

    width = attrs.get("width")
    width_pt = Pt(width) if width else None

    doc.add_picture(io.BytesIO(img_data), width=width_pt)


# ---------------------------------------------------------------------------
# Table
# ---------------------------------------------------------------------------

def _render_table(doc: Document, node: dict, images: dict[str, bytes]) -> None:
    rows_data = node.get("content", [])
    if not rows_data:
        return

    num_rows = len(rows_data)
    num_cols = max(len(r.get("content", [])) for r in rows_data) if rows_data else 0
    if num_cols == 0:
        return

    table = doc.add_table(rows=num_rows, cols=num_cols, style="Table Grid")

    for row_idx, row_node in enumerate(rows_data):
        cells = row_node.get("content", [])
        for col_idx, cell_node in enumerate(cells):
            if col_idx >= num_cols:
                break
            cell = table.cell(row_idx, col_idx)
            cell_attrs = cell_node.get("attrs", {})

            bg_color = cell_attrs.get("backgroundColor")
            if bg_color:
                _set_cell_shading(cell, bg_color.lstrip("#"))

            cell_content = cell_node.get("content", [])
            if cell_content:
                # Clear default empty paragraph
                cell.text = ""
                for i, child in enumerate(cell_content):
                    if i == 0 and cell.paragraphs:
                        para = cell.paragraphs[0]
                    else:
                        para = cell.add_paragraph()
                    child_attrs = child.get("attrs", {})
                    _apply_alignment(para, child_attrs)
                    _apply_inline_content(para, child.get("content", []), images)


def _set_cell_shading(cell, hex_color: str) -> None:
    tc_pr = cell._element.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), hex_color)
    tc_pr.append(shading)


# ---------------------------------------------------------------------------
# Lists
# ---------------------------------------------------------------------------

def _render_list(
    doc: Document, node: dict, images: dict[str, bytes], style: str
) -> None:
    for item_node in node.get("content", []):
        if item_node.get("type") != "listItem":
            continue
        for child in item_node.get("content", []):
            para = doc.add_paragraph(style=style)
            child_attrs = child.get("attrs", {})
            _apply_alignment(para, child_attrs)
            _apply_inline_content(para, child.get("content", []), images)


# ---------------------------------------------------------------------------
# Blockquote
# ---------------------------------------------------------------------------

def _render_blockquote(doc: Document, node: dict, images: dict[str, bytes]) -> None:
    for child in node.get("content", []):
        try:
            para = doc.add_paragraph(style="Quote")
        except KeyError:
            para = doc.add_paragraph()
        child_attrs = child.get("attrs", {})
        _apply_alignment(para, child_attrs)
        _apply_inline_content(para, child.get("content", []), images)


# ---------------------------------------------------------------------------
# Page break
# ---------------------------------------------------------------------------

def _render_page_break(doc: Document) -> None:
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(WD_BREAK.PAGE)


# ---------------------------------------------------------------------------
# Horizontal rule
# ---------------------------------------------------------------------------

def _render_horizontal_rule(doc: Document) -> None:
    para = doc.add_paragraph()
    p_pr = para._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


# ---------------------------------------------------------------------------
# Inline content (text, marks, nested images)
# ---------------------------------------------------------------------------

def _apply_inline_content(
    para, content: list[dict], images: dict[str, bytes]
) -> None:
    for item in content:
        item_type = item.get("type", "")
        if item_type == "text":
            text = item.get("text", "")
            run = para.add_run(text)
            _apply_marks(run, item.get("marks", []))
        elif item_type == "image":
            src = item.get("attrs", {}).get("src", "")
            img_data = images.get(src)
            if img_data:
                width = item.get("attrs", {}).get("width")
                run = para.add_run()
                run.add_picture(io.BytesIO(img_data), width=Pt(width) if width else None)
        elif item_type == "hardBreak":
            run = para.add_run()
            run.add_break()


def _apply_marks(run, marks: list[dict]) -> None:
    for mark in marks:
        mark_type = mark.get("type", "")
        if mark_type == "bold":
            run.bold = True
        elif mark_type == "italic":
            run.italic = True
        elif mark_type == "underline":
            run.underline = True
        elif mark_type == "strike":
            run.font.strike = True
        elif mark_type == "textStyle":
            attrs = mark.get("attrs") or {}
            font_family = attrs.get("fontFamily")
            if font_family:
                run.font.name = font_family
            font_size = attrs.get("fontSize")
            if font_size:
                try:
                    pts = float(str(font_size).replace("pt", "").strip())
                    run.font.size = Pt(pts)
                except (ValueError, AttributeError):
                    pass
            color = attrs.get("color")
            if color:
                try:
                    run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
                except (ValueError, AttributeError):
                    pass


# ---------------------------------------------------------------------------
# Alignment helper
# ---------------------------------------------------------------------------

_ALIGN_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}


def _apply_alignment(para, attrs: dict[str, Any]) -> None:
    align = attrs.get("textAlign")
    if align and align in _ALIGN_MAP:
        para.alignment = _ALIGN_MAP[align]
