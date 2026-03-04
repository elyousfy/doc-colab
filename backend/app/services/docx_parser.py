"""Parse a .docx file into Tiptap-compatible JSON using python-docx (no HTML middleman)."""

from __future__ import annotations

import io
import re
import uuid
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

EMU_TO_PT = 72.0 / 914400.0

_MIME_EXT = {
    "image/png": ".png",
    "image/jpeg": ".jpg",
    "image/gif": ".gif",
    "image/bmp": ".bmp",
    "image/tiff": ".tiff",
    "image/svg+xml": ".svg",
    "image/x-emf": ".emf",
    "image/x-wmf": ".wmf",
}


def _block_id() -> str:
    return str(uuid.uuid4())


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def parse_docx(file_bytes: bytes) -> tuple[dict, list[dict]]:
    """Parse a DOCX file into (tiptap_json_content, images_list).

    images_list items: {"filename": str, "data": bytes, "mime_type": str}
    """
    doc = Document(io.BytesIO(file_bytes))
    images: list[dict] = []
    image_rels = _extract_image_rels(doc)

    # --- Header / Footer ---
    header_blocks = _parse_section_headers(doc, images)
    footer_blocks = _parse_section_footers(doc, images)

    # --- Body ---
    para_map = {p._element: p for p in doc.paragraphs}
    table_map = {t._element: t for t in doc.tables}

    body_blocks: list[dict] = []
    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p" and element in para_map:
            node = _parse_paragraph(para_map[element], image_rels, images)
            if node is not None:
                body_blocks.append(node)
        elif tag == "p":
            body_blocks.append({"type": "paragraph", "attrs": {"blockId": _block_id()}})

        elif tag == "tbl" and element in table_map:
            body_blocks.append(_parse_table(table_map[element], image_rels, images))

        elif tag == "sdt":
            sdt_nodes = _parse_sdt(element, image_rels, images, para_map, table_map)
            body_blocks.extend(sdt_nodes)

    body_blocks = post_process_lists(body_blocks)

    # --- Assemble final doc ---
    blocks: list[dict] = []
    if header_blocks:
        blocks.append({
            "type": "docSection",
            "attrs": {"blockId": _block_id(), "sectionType": "header"},
            "content": header_blocks,
        })
    blocks.extend(body_blocks)
    if footer_blocks:
        blocks.append({
            "type": "docSection",
            "attrs": {"blockId": _block_id(), "sectionType": "footer"},
            "content": footer_blocks,
        })

    content = {"type": "doc", "content": blocks}
    return content, images


# ---------------------------------------------------------------------------
# Header / Footer extraction
# ---------------------------------------------------------------------------

def _parse_section_headers(doc: Document, images: list[dict]) -> list[dict]:
    """Extract header content from document sections."""
    for section in doc.sections:
        header = section.header
        if header.is_linked_to_previous:
            continue
        blocks = _parse_header_footer_part(header, images)
        if blocks:
            return blocks
    # Fallback: try first section even if linked
    if doc.sections:
        blocks = _parse_header_footer_part(doc.sections[0].header, images)
        if blocks:
            return blocks
    return []


def _parse_section_footers(doc: Document, images: list[dict]) -> list[dict]:
    """Extract footer content from document sections."""
    for section in doc.sections:
        footer = section.footer
        if footer.is_linked_to_previous:
            continue
        blocks = _parse_header_footer_part(footer, images)
        if blocks:
            return blocks
    if doc.sections:
        blocks = _parse_header_footer_part(doc.sections[0].footer, images)
        if blocks:
            return blocks
    return []


def _parse_header_footer_part(part_obj, images: list[dict]) -> list[dict]:
    """Parse paragraphs and images from a header or footer object."""
    hf_image_rels = _extract_image_rels_from_part(part_obj)
    blocks: list[dict] = []
    for para in part_obj.paragraphs:
        text = para.text.strip()
        inline_imgs = _extract_inline_images(para, hf_image_rels, images)
        if not text and not inline_imgs:
            continue
        node = _parse_paragraph(para, hf_image_rels, images)
        if node is not None:
            blocks.append(node)
    # Also check for tables in headers/footers
    for table in part_obj.tables:
        blocks.append(_parse_table(table, hf_image_rels, images))
    return blocks


def _extract_image_rels_from_part(part_obj) -> dict[str, dict]:
    """Extract image rels from a header/footer part (which has its own rels)."""
    rels: dict[str, dict] = {}
    try:
        for rel_id, rel in part_obj.part.rels.items():
            if "image" in rel.reltype:
                image_part = rel.target_part
                ct = image_part.content_type
                ext = _MIME_EXT.get(ct, ".bin")
                fname = f"img_{rel_id}_hf{ext}"
                rels[rel_id] = {
                    "filename": fname,
                    "data": image_part.blob,
                    "mime_type": ct,
                }
    except Exception:
        pass
    return rels


# ---------------------------------------------------------------------------
# SDT (Structured Document Tag) parsing — handles Table of Contents, etc.
# ---------------------------------------------------------------------------

def _parse_sdt(
    sdt_element,
    image_rels: dict[str, dict],
    images: list[dict],
    para_map: dict,
    table_map: dict,
) -> list[dict]:
    """Parse a w:sdt element, extracting its content paragraphs/tables.

    If it's a Table of Contents, wraps in a docSection with sectionType='toc'.
    Otherwise, returns content as flat blocks.
    """
    sdt_content = sdt_element.find(qn("w:sdtContent"))
    if sdt_content is None:
        return []

    is_toc = _sdt_is_toc(sdt_element)

    inner_blocks: list[dict] = []
    for child in sdt_content:
        child_tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

        if child_tag == "p":
            if child in para_map:
                node = _parse_paragraph(para_map[child], image_rels, images)
            else:
                # Paragraph not in doc.paragraphs (SDT-owned) — parse raw
                node = _parse_sdt_paragraph(child, image_rels, images)
            if node is not None:
                inner_blocks.append(node)

        elif child_tag == "tbl":
            if child in table_map:
                inner_blocks.append(_parse_table(table_map[child], image_rels, images))

        elif child_tag == "sdt":
            inner_blocks.extend(_parse_sdt(child, image_rels, images, para_map, table_map))

    if not inner_blocks:
        return []

    if is_toc:
        return [{
            "type": "docSection",
            "attrs": {"blockId": _block_id(), "sectionType": "toc"},
            "content": inner_blocks,
        }]
    return inner_blocks


def _sdt_is_toc(sdt_element) -> bool:
    """Detect if an SDT element is a Table of Contents."""
    sdt_pr = sdt_element.find(qn("w:sdtPr"))
    if sdt_pr is None:
        return False
    doc_part = sdt_pr.find(qn("w:docPartObj"))
    if doc_part is not None:
        gallery = doc_part.find(qn("w:docPartGallery"))
        if gallery is not None:
            val = gallery.get(qn("w:val"), "")
            if "Table of Contents" in val or "TOC" in val.upper():
                return True
    # Also check for alias
    alias = sdt_pr.find(qn("w:alias"))
    if alias is not None:
        val = alias.get(qn("w:val"), "")
        if "toc" in val.lower() or "table of contents" in val.lower() or "content" in val.lower():
            return True
    return False


def _parse_sdt_paragraph(
    p_element,
    image_rels: dict[str, dict],
    images: list[dict],
) -> dict | None:
    """Parse a raw paragraph XML element that isn't in doc.paragraphs (SDT-owned)."""
    try:
        para = Paragraph(p_element, None)
    except Exception:
        para = Paragraph(p_element, p_element.getparent())

    text_content = _parse_runs(para.runs)
    inline_images = _extract_inline_images_raw(p_element, image_rels, images)
    if inline_images:
        text_content = inline_images + text_content

    if not text_content:
        return None

    attrs: dict[str, Any] = {"blockId": _block_id()}
    pPr = p_element.find(qn("w:pPr"))
    if pPr is not None:
        jc = pPr.find(qn("w:jc"))
        if jc is not None:
            val = jc.get(qn("w:val"), "")
            mapping = {"left": "left", "center": "center", "right": "right", "both": "justify"}
            if val in mapping:
                attrs["textAlign"] = mapping[val]

    # Detect TOC style → render as paragraph (TOC entries are usually styled TOC1, TOC2, etc.)
    style_el = pPr.find(qn("w:pStyle")) if pPr is not None else None
    style_name = style_el.get(qn("w:val"), "") if style_el is not None else ""
    heading_match = _HEADING_RE.search(style_name)

    if heading_match:
        level = int(heading_match.group(1))
        level = max(1, min(level, 6))
        return {"type": "heading", "attrs": {**attrs, "level": level}, "content": text_content}

    return {"type": "paragraph", "attrs": attrs, "content": text_content}


def _extract_inline_images_raw(
    p_element,
    image_rels: dict[str, dict],
    images: list[dict],
) -> list[dict]:
    """Extract inline images from a raw paragraph element."""
    nodes: list[dict] = []
    for drawing in p_element.findall(f".//{qn('w:drawing')}"):
        blip = drawing.find(f".//{qn('a:blip')}")
        if blip is None:
            continue
        r_embed = blip.get(qn("r:embed"))
        if not r_embed or r_embed not in image_rels:
            continue
        img_info = image_rels[r_embed]
        if not any(i["filename"] == img_info["filename"] for i in images):
            images.append(img_info)
        width_pt, height_pt = _get_drawing_dimensions(drawing)
        img_node: dict = {
            "type": "image",
            "attrs": {
                "src": f"__IMAGE__{img_info['filename']}",
                "alt": "",
                "title": None,
                "blockId": _block_id(),
            },
        }
        if width_pt:
            img_node["attrs"]["width"] = round(width_pt, 1)
        if height_pt:
            img_node["attrs"]["height"] = round(height_pt, 1)
        nodes.append(img_node)
    return nodes


# ---------------------------------------------------------------------------
# Image relationship extraction
# ---------------------------------------------------------------------------

def _extract_image_rels(doc: Document) -> dict[str, dict]:
    """Build rId → {data, mime_type, filename} for all images in the document part."""
    rels: dict[str, dict] = {}
    for rel_id, rel in doc.part.rels.items():
        if "image" in rel.reltype:
            image_part = rel.target_part
            ct = image_part.content_type
            ext = _MIME_EXT.get(ct, ".bin")
            fname = f"img_{rel_id}{ext}"
            rels[rel_id] = {
                "filename": fname,
                "data": image_part.blob,
                "mime_type": ct,
            }
    return rels


# ---------------------------------------------------------------------------
# Paragraph parsing
# ---------------------------------------------------------------------------

_HEADING_RE = re.compile(r"[Hh]eading\s*(\d)", re.IGNORECASE)
_LIST_STYLES = {
    "list bullet": "bullet",
    "list number": "ordered",
    "list paragraph": "bullet",
}


def _parse_paragraph(
    para: Paragraph,
    image_rels: dict[str, dict],
    images: list[dict],
) -> dict | None:
    """Convert a single paragraph to a Tiptap node."""

    # Check for images inline
    inline_images = _extract_inline_images(para, image_rels, images)

    # Detect heading from style
    style_name = (para.style.name or "").strip() if para.style else ""
    heading_match = _HEADING_RE.search(style_name)

    # Detect list from numPr XML element
    pPr = para._element.find(qn("w:pPr"))
    num_pr = pPr.find(qn("w:numPr")) if pPr is not None else None
    list_type = None
    list_level = 0
    if num_pr is not None:
        ilvl_el = num_pr.find(qn("w:ilvl"))
        list_level = int(ilvl_el.get(qn("w:val"), "0")) if ilvl_el is not None else 0
        num_id_el = num_pr.find(qn("w:numId"))
        num_id = num_id_el.get(qn("w:val"), "0") if num_id_el is not None else "0"
        list_type = "ordered" if int(num_id) % 2 == 0 else "bullet"
    else:
        lower_name = style_name.lower()
        for key, ltype in _LIST_STYLES.items():
            if key in lower_name:
                list_type = ltype
                break

    # Build text content from runs
    text_content = _parse_runs(para.runs)

    # Merge inline images into content
    if inline_images:
        text_content = inline_images + text_content

    # Paragraph attributes
    attrs: dict[str, Any] = {"blockId": _block_id()}
    alignment = _get_alignment(para)
    if alignment:
        attrs["textAlign"] = alignment

    ppr = para._element.find(qn("w:pPr"))
    if ppr is not None:
        rpr = ppr.find(qn("w:rPr"))
        if rpr is not None:
            font_el = rpr.find(qn("w:rFonts"))
            if font_el is not None:
                fname = font_el.get(qn("w:ascii")) or font_el.get(qn("w:hAnsi"))
                if fname:
                    attrs["fontFamily"] = fname
            sz_el = rpr.find(qn("w:sz"))
            if sz_el is not None:
                half_points = int(sz_el.get(qn("w:val"), "0"))
                if half_points:
                    attrs["fontSize"] = f"{half_points // 2}pt"

    if heading_match:
        level = int(heading_match.group(1))
        level = max(1, min(level, 6))
        node: dict = {
            "type": "heading",
            "attrs": {**attrs, "level": level},
        }
        if text_content:
            node["content"] = text_content
        return node

    if list_type:
        node = {
            "type": "paragraph",
            "attrs": attrs,
            "_isList": True,
            "_listType": list_type,
            "_listLevel": list_level,
        }
        if text_content:
            node["content"] = text_content
        return node

    # Check for page break
    for run in para.runs:
        for br in run._element.findall(qn("w:br")):
            if br.get(qn("w:type")) == "page":
                return {
                    "type": "hardBreak",
                    "attrs": {"blockId": _block_id(), "pageBreak": True},
                }

    node = {"type": "paragraph", "attrs": attrs}
    if text_content:
        node["content"] = text_content
    return node


def _get_alignment(para: Paragraph) -> str | None:
    ppr = para._element.find(qn("w:pPr"))
    if ppr is not None:
        jc = ppr.find(qn("w:jc"))
        if jc is not None:
            val = jc.get(qn("w:val"), "")
            mapping = {"left": "left", "center": "center", "right": "right", "both": "justify"}
            return mapping.get(val)
    return None


# ---------------------------------------------------------------------------
# Run parsing
# ---------------------------------------------------------------------------

def _parse_runs(runs) -> list[dict]:
    """Convert paragraph runs into Tiptap text/mark nodes."""
    nodes: list[dict] = []
    for run in runs:
        text = run.text
        if not text:
            continue
        marks = _get_run_marks(run)
        node: dict = {"type": "text", "text": text}
        if marks:
            node["marks"] = marks
        nodes.append(node)
    return nodes


def _get_run_marks(run) -> list[dict]:
    marks: list[dict] = []
    if run.bold:
        marks.append({"type": "bold"})
    if run.italic:
        marks.append({"type": "italic"})
    if run.underline:
        marks.append({"type": "underline"})
    if run.font.strike:
        marks.append({"type": "strike"})

    style_attrs: dict[str, Any] = {}
    if run.font.name:
        style_attrs["fontFamily"] = run.font.name
    if run.font.size:
        pt_size = run.font.size.pt
        style_attrs["fontSize"] = f"{pt_size}pt"
    if run.font.color and run.font.color.rgb:
        style_attrs["color"] = f"#{run.font.color.rgb}"

    if style_attrs:
        marks.append({"type": "textStyle", "attrs": style_attrs})

    return marks


# ---------------------------------------------------------------------------
# Inline image extraction
# ---------------------------------------------------------------------------

def _extract_inline_images(
    para: Paragraph,
    image_rels: dict[str, dict],
    images: list[dict],
) -> list[dict]:
    """Extract inline images (<w:drawing>) from a paragraph element."""
    nodes: list[dict] = []
    for drawing in para._element.findall(f".//{qn('w:drawing')}"):
        blip = drawing.find(f".//{qn('a:blip')}")
        if blip is None:
            continue
        r_embed = blip.get(qn("r:embed"))
        if not r_embed or r_embed not in image_rels:
            continue

        img_info = image_rels[r_embed]
        if not any(i["filename"] == img_info["filename"] for i in images):
            images.append(img_info)

        width_pt, height_pt = _get_drawing_dimensions(drawing)

        img_node: dict = {
            "type": "image",
            "attrs": {
                "src": f"__IMAGE__{img_info['filename']}",
                "alt": "",
                "title": None,
                "blockId": _block_id(),
            },
        }
        if width_pt:
            img_node["attrs"]["width"] = round(width_pt, 1)
        if height_pt:
            img_node["attrs"]["height"] = round(height_pt, 1)
        nodes.append(img_node)

    return nodes


def _get_drawing_dimensions(drawing) -> tuple[float | None, float | None]:
    extent = drawing.find(f".//{qn('wp:extent')}")
    if extent is None:
        return None, None
    cx = extent.get("cx")
    cy = extent.get("cy")
    w = int(cx) * EMU_TO_PT if cx else None
    h = int(cy) * EMU_TO_PT if cy else None
    return w, h


# ---------------------------------------------------------------------------
# Table parsing
# ---------------------------------------------------------------------------

def _parse_table(
    table: Table,
    image_rels: dict[str, dict],
    images: list[dict],
) -> dict:
    rows: list[dict] = []
    for row in table.rows:
        cells: list[dict] = []
        for cell in row.cells:
            cell_content: list[dict] = []
            for para in cell.paragraphs:
                node = _parse_paragraph(para, image_rels, images)
                if node is not None:
                    node.pop("_isList", None)
                    node.pop("_listType", None)
                    node.pop("_listLevel", None)
                    cell_content.append(node)

            cell_attrs: dict[str, Any] = {}
            tc_pr = cell._element.find(qn("w:tcPr"))
            if tc_pr is not None:
                shd = tc_pr.find(qn("w:shd"))
                if shd is not None:
                    fill = shd.get(qn("w:fill"))
                    if fill and fill != "auto":
                        cell_attrs["backgroundColor"] = f"#{fill}"
                tc_w = tc_pr.find(qn("w:tcW"))
                if tc_w is not None:
                    w_val = tc_w.get(qn("w:w"))
                    w_type = tc_w.get(qn("w:type"))
                    if w_val and w_type == "dxa":
                        cell_attrs["width"] = f"{int(w_val) / 20}pt"

            cell_node: dict = {
                "type": "tableCell",
                "attrs": cell_attrs,
            }
            if cell_content:
                cell_node["content"] = cell_content
            else:
                cell_node["content"] = [{"type": "paragraph", "attrs": {"blockId": _block_id()}}]
            cells.append(cell_node)

        rows.append({"type": "tableRow", "content": cells})

    return {
        "type": "table",
        "attrs": {"blockId": _block_id()},
        "content": rows,
    }


# ---------------------------------------------------------------------------
# List post-processing
# ---------------------------------------------------------------------------

def post_process_lists(blocks: list[dict]) -> list[dict]:
    """Group consecutive _isList paragraphs into bulletList/orderedList wrappers."""
    result: list[dict] = []
    i = 0
    while i < len(blocks):
        block = blocks[i]
        if block.get("_isList"):
            list_type = block.get("_listType", "bullet")
            wrapper_type = "bulletList" if list_type == "bullet" else "orderedList"
            items: list[dict] = []

            while i < len(blocks) and blocks[i].get("_isList") and blocks[i].get("_listType") == list_type:
                item_block = blocks[i]
                item_block.pop("_isList", None)
                item_block.pop("_listType", None)
                item_block.pop("_listLevel", None)
                items.append({
                    "type": "listItem",
                    "content": [item_block],
                })
                i += 1

            result.append({
                "type": wrapper_type,
                "attrs": {"blockId": _block_id()},
                "content": items,
            })
        else:
            result.append(block)
            i += 1

    return result
